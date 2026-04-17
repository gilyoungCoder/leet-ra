"""
LEET-RA docx 출력기 (한컴오피스에서 hwp로 저장 가능한 안정 포맷).

Usage:
    python export_docx.py --input samples/v2_5questions.json --output output/문제지.docx
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def load_questions(src: Path) -> list[dict[str, Any]]:
    data = json.loads(src.read_text(encoding="utf-8"))
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        for key in ("questions", "items", "results"):
            if key in data and isinstance(data[key], list):
                return data[key]
        return [data]
    raise ValueError(f"unsupported JSON shape: {type(data).__name__}")


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 10, indent: float = 0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "함초롬바탕"
    return p


def write_docx(out_path: Path, title: str, questions: list[dict[str, Any]]) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # A4, 좁은 여백, 2단 구성
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    # 2단 레이아웃
    sectPr = section._sectPr
    from docx.oxml.ns import qn
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        from docx.oxml import OxmlElement
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:sep"), "1")
    cols.set(qn("w:space"), "720")

    add_para(doc, title, bold=True, size=16)
    add_para(doc, f"총 {len(questions)}문항", size=9)
    add_para(doc, "")

    for q in questions:
        qno = q.get("question_number", "")
        stem = q.get("stem", "")
        passage = q.get("passage_text", "")
        bogie = q.get("bogie_items", {}) or {}
        choices = q.get("choices", []) or []
        answer = q.get("answer", "")
        explanations = q.get("explanations", {}) or {}
        domain = q.get("domain", "")

        add_para(doc, f"{qno}. [{domain}] {stem}", bold=True, size=11)
        if passage:
            add_para(doc, "<지문>", bold=True)
            for line in str(passage).splitlines():
                line = line.strip()
                if line:
                    add_para(doc, line)
        if bogie:
            add_para(doc, "<보기>", bold=True)
            for k, v in bogie.items():
                add_para(doc, f"{k}. {v}", indent=0.5)
        for c in choices:
            add_para(doc, str(c))
        if answer:
            add_para(doc, f"[정답] {answer}", bold=True)
        if explanations:
            add_para(doc, "[해설]", bold=True)
            for k, v in explanations.items():
                add_para(doc, f"{k}) {v}", size=9, indent=0.5)
        add_para(doc, "")

    doc.save(out_path)
    return out_path


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", "-i", required=True)
    ap.add_argument("--output", "-o", required=True)
    ap.add_argument("--title", "-t", default="LEET 추리논증 문제지")
    args = ap.parse_args()
    src = Path(args.input)
    dst = Path(args.output)
    qs = load_questions(src)
    write_docx(dst, args.title, qs)
    print(f"[ok] wrote {dst}  ({len(qs)} 문항)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
